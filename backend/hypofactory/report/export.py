"""Экспорт результата анализа в бизнес-форматы.

DOCX/PDF/CSV — отчёты; tasks-CSV — задачи для импорта в Jira/YouTrack;
protocol-DOCX — протоколы лабораторной проверки гипотез (с полями под результаты).
На вход подаётся результат /api/analyze (dict). Все функции возвращают bytes.
PDF использует вшитый шрифт DejaVuSans (кириллица гарантирована в любом окружении).
"""
from __future__ import annotations

import csv
import io
import os


# ---------------- DOCX ----------------

def to_docx(result: dict) -> bytes:
    import docx
    from docx.shared import Pt

    d = docx.Document()
    rep = result.get("report", {})
    d.add_heading("Фабрика гипотез — отчёт", level=0)
    d.add_paragraph(f"Источник данных: {rep.get('source_file', '—')}")

    d.add_heading("Сводка по хвостам", level=1)
    t = d.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in [
        ("Переработка, СМТ", rep.get("feed_smt")),
        ("Хвосты, СМТ", rep.get("tails_smt")),
        ("Потери Элемент 28, т", rep.get("tails_loss_t_28")),
        ("Потери Элемент 29, т", rep.get("tails_loss_t_29")),
    ]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = f"{v:,.0f}".replace(",", " ") if isinstance(v, (int, float)) else "—"

    d.add_heading("Диагноз", level=1)
    d.add_paragraph(result.get("dominant_strategy", ""))
    goal = (result.get("meta") or {}).get("user_goal")
    if goal:
        d.add_paragraph("Целевая задача и ограничения технолога: " + goal)

    d.add_heading("Гипотезы", level=1)
    for h in result.get("hypotheses", []):
        d.add_heading(f"{h.get('id')}. {h.get('title')}", level=2)
        d.add_paragraph(h.get("statement", ""))
        p = d.add_paragraph()
        ev = h.get("expected_value", {}) or {}
        nov = h.get("novelty", {}) or {}
        run = p.add_run(
            f"Ценность: {ev.get('addressable_recoverable_t', 0):,.0f} т извлекаемых потерь · "
            f"Новизна: {nov.get('label', '—')} ({nov.get('score', '')}) · "
            f"Категория: {(h.get('target') or {}).get('category', '')}".replace(",", " "))
        run.font.size = Pt(9)
        d.add_paragraph("Обоснование: " + h.get("rationale", ""))
        d.add_paragraph("Механизм влияния: " + h.get("mechanism_of_influence", ""))
        if ev.get("kpi_text"):
            d.add_paragraph("Ожидаемый эффект (KPI): " + ev["kpi_text"])
        if ev.get("value_usd"):
            d.add_paragraph(f"Экономическая оценка: адресуемые потери ≈ ${_n(ev['value_usd'])}/год "
                            f"(референсные цены); возврат 10% ≈ ${_n(ev['value_usd'] * 0.1)}/год, "
                            f"25% ≈ ${_n(ev['value_usd'] * 0.25)}/год.")
        srcs = h.get("sources", [])
        if srcs:
            d.add_paragraph("Источники:")
            for s in srcs:
                page = f", с.{s['page']}" if s.get("page") else ""
                d.add_paragraph(f"— {s.get('title','')}{page}: «{s.get('quote','')[:180]}»",
                                style="List Bullet")
        risks = h.get("risks", {}) or {}
        if risks.get("technical") or risks.get("economic"):
            d.add_paragraph("Риски:")
            for r in risks.get("technical", []):
                d.add_paragraph(f"— (тех.) {r}", style="List Bullet")
            for r in risks.get("economic", []):
                d.add_paragraph(f"— (экон.) {r}", style="List Bullet")
        road = h.get("verification_roadmap", [])
        if road:
            d.add_paragraph("Дорожная карта проверки:")
            for i, v in enumerate(road, 1):
                d.add_paragraph(
                    f"{i}. {v.get('step','')} | Ресурсы: {v.get('resources','—')} | "
                    f"Критерий: {v.get('success_criteria','—')}", style="List Number")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


# ---------------- PDF ----------------

_FONT = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")


def to_pdf(result: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable)
    from reportlab.lib import colors

    if "DejaVu" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVu", _FONT))

    def st(size=10, bold=False, color="#111111", space=4):
        return ParagraphStyle("s", fontName="DejaVu", fontSize=size, leading=size * 1.35,
                              textColor=color, spaceAfter=space)

    rep = result.get("report", {})
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm,
                            topMargin=16 * mm, bottomMargin=16 * mm,
                            title="Фабрика гипотез — отчёт")
    story = [
        Paragraph("Фабрика гипотез — отчёт", st(18, space=2)),
        Paragraph(f"Источник данных: {rep.get('source_file','—')}", st(10, color="#555555", space=10)),
    ]

    rows = [["Переработка, СМТ", "Хвосты, СМТ", "Потери Эл.28, т", "Потери Эл.29, т"],
            [_n(rep.get("feed_smt")), _n(rep.get("tails_smt")),
             _n(rep.get("tails_loss_t_28")), _n(rep.get("tails_loss_t_29"))]]
    tbl = Table(rows, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), "DejaVu"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef2f7")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#c9d3de")),
        ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story += [tbl, Spacer(1, 8),
              Paragraph("<b>Диагноз.</b> " + result.get("dominant_strategy", ""), st(10.5, space=12))]
    _goal = (result.get("meta") or {}).get("user_goal")
    if _goal:
        story.append(Paragraph("<b>Целевая задача и ограничения.</b> " + _goal, st(10.5, space=12)))

    for h in result.get("hypotheses", []):
        ev = h.get("expected_value", {}) or {}
        nov = h.get("novelty", {}) or {}
        story.append(HRFlowable(width="100%", color=colors.HexColor("#c9d3de"), spaceAfter=6))
        story.append(Paragraph(f"{h.get('id')}. {h.get('title')}", st(12.5, space=3)))
        story.append(Paragraph(h.get("statement", ""), st(10, space=3)))
        story.append(Paragraph(
            f"Ценность: {_n(ev.get('addressable_recoverable_t'))} т · Новизна: {nov.get('label','—')} "
            f"({nov.get('score','')}) · {(h.get('target') or {}).get('category','')}",
            st(8.5, color="#666666", space=5)))
        story.append(Paragraph("<b>Обоснование:</b> " + h.get("rationale", ""), st(9.5)))
        story.append(Paragraph("<b>Механизм:</b> " + h.get("mechanism_of_influence", ""), st(9.5)))
        if ev.get("kpi_text"):
            story.append(Paragraph("<b>Эффект (KPI):</b> " + ev["kpi_text"], st(9.5)))
        if ev.get("value_usd"):
            story.append(Paragraph(
                f"<b>Экономика:</b> адресуемые потери ≈ ${_n(ev['value_usd'])}/год; "
                f"возврат 10% ≈ ${_n(ev['value_usd'] * 0.1)}, 25% ≈ ${_n(ev['value_usd'] * 0.25)} (реф. цены)",
                st(9.5)))
        for s in h.get("sources", []):
            page = f", с.{s['page']}" if s.get("page") else ""
            story.append(Paragraph(f"Источник: {s.get('title','')}{page} — «{s.get('quote','')[:160]}»",
                                   st(8.5, color="#555555")))
        risks = h.get("risks", {}) or {}
        rk = "; ".join(risks.get("technical", []) + risks.get("economic", []))
        if rk:
            story.append(Paragraph("<b>Риски:</b> " + rk, st(9)))
        road = h.get("verification_roadmap", [])
        for i, v in enumerate(road, 1):
            story.append(Paragraph(
                f"{i}) {v.get('step','')} — критерий: {v.get('success_criteria','—')}", st(9)))
        story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()


def _n(v) -> str:
    return f"{v:,.0f}".replace(",", " ") if isinstance(v, (int, float)) else "—"


# ---------------- CSV ----------------

def to_csv(result: dict) -> bytes:
    buf = io.StringIO()
    w = csv.writer(buf, delimiter=";")
    w.writerow(["id", "title", "statement", "rationale", "mechanism_of_influence",
                "stream", "size_classes", "category",
                "value_t", "value_usd", "kpi_text", "novelty", "novelty_score",
                "risks_technical", "risks_economic", "confidence", "priority_score", "sources"])
    for h in result.get("hypotheses", []):
        t = h.get("target", {}) or {}
        ev = h.get("expected_value", {}) or {}
        nov = h.get("novelty", {}) or {}
        rk = h.get("risks", {}) or {}
        srcs = " | ".join(f"{s.get('title','')} с.{s.get('page','')}" for s in h.get("sources", []))
        w.writerow([h.get("id"), h.get("title"), h.get("statement"), h.get("rationale"),
                    h.get("mechanism_of_influence"), t.get("stream"),
                    ",".join(t.get("size_classes", [])), t.get("category"),
                    ev.get("addressable_recoverable_t"), ev.get("value_usd"), ev.get("kpi_text"),
                    nov.get("label"), nov.get("score"),
                    " | ".join(rk.get("technical", [])), " | ".join(rk.get("economic", [])),
                    h.get("confidence"), h.get("priority_score"), srcs])
    return ("﻿" + buf.getvalue()).encode("utf-8")  # BOM для Excel


# ---------------- Задачи для Jira / YouTrack ----------------

def to_tasks_csv(result: dict) -> bytes:
    """CSV в формате импорта задач Jira/YouTrack (Summary/Description/Priority/Labels).

    Jira: Project settings → Import CSV; YouTrack: Import → CSV. Разделитель — запятая.
    """
    def _priority(tons) -> str:
        if not isinstance(tons, (int, float)):
            return "Medium"
        return "High" if tons >= 1000 else "Medium" if tons >= 100 else "Low"

    def _label(s: str) -> str:
        return (s or "").strip().lower().replace(" ", "-").replace(",", "")

    src = (result.get("report") or {}).get("source_file", "")
    buf = io.StringIO()
    w = csv.writer(buf)  # запятая — дефолт импорта Jira/YouTrack
    w.writerow(["Summary", "Description", "Issue Type", "Priority", "Labels"])
    for h in result.get("hypotheses", []):
        t = h.get("target", {}) or {}
        ev = h.get("expected_value", {}) or {}
        road = h.get("verification_roadmap", [])
        desc_lines = [
            h.get("statement", ""),
            "",
            f"Обоснование: {h.get('rationale', '')}",
            f"Механизм влияния: {h.get('mechanism_of_influence', '')}",
            f"Ожидаемая ценность: {ev.get('kpi_text', '') or str(ev.get('addressable_recoverable_t', '')) + ' т извлекаемых потерь'}",
            (f"Экономическая оценка: адресуемые потери ≈ ${ev['value_usd']:,.0f}/год (референсные цены)".replace(",", " ")
             if ev.get("value_usd") else ""),
            f"Оценка новизны: {(h.get('novelty') or {}).get('label', '—')} "
            f"({(h.get('novelty') or {}).get('score', '')})",
            "Риски: " + ("; ".join((h.get("risks") or {}).get("technical", [])
                                   + (h.get("risks") or {}).get("economic", [])) or "—"),
            f"Источник данных: {src}",
        ]
        if road:
            desc_lines.append("")
            desc_lines.append("План проверки:")
            for i, v in enumerate(road, 1):
                desc_lines.append(f"{i}. {v.get('step', '')} | Ресурсы: {v.get('resources', '—')} | "
                                  f"Критерий успеха: {v.get('success_criteria', '—')}")
        srcs = h.get("sources", [])
        if srcs:
            desc_lines.append("")
            desc_lines.append("Литература: " + "; ".join(
                f"{s.get('title', '')}{', с.' + str(s['page']) if s.get('page') else ''}" for s in srcs))
        labels = " ".join(filter(None, [
            "hypothesis-factory", _label(t.get("stream", "")), _label(t.get("mechanism", ""))]))
        w.writerow([f"{h.get('id')}. {h.get('title', '')}", "\n".join(desc_lines),
                    "Task", _priority(ev.get("addressable_recoverable_t")), labels])
    return ("﻿" + buf.getvalue()).encode("utf-8")


# ---------------- Протоколы лабораторной проверки ----------------

def to_protocol_docx(result: dict) -> bytes:
    """Протокол эксперимента на каждую гипотезу: методика + поля под результаты."""
    import docx
    from docx.enum.text import WD_BREAK

    rep = result.get("report", {})
    d = docx.Document()
    d.add_heading("Протоколы лабораторной проверки гипотез", level=0)
    d.add_paragraph(f"Источник данных: {rep.get('source_file', '—')}")
    d.add_paragraph(f"Диагноз: {result.get('dominant_strategy', '')}")

    hyps = result.get("hypotheses", [])
    for n, h in enumerate(hyps):
        ev = h.get("expected_value", {}) or {}
        t = h.get("target", {}) or {}
        d.add_heading(f"Протокол {h.get('id')}: {h.get('title', '')}", level=1)

        tab = d.add_table(rows=0, cols=2)
        tab.style = "Light Grid Accent 1"
        for k, v in [
            ("Гипотеза", h.get("statement", "")),
            ("Обоснование", h.get("rationale", "")),
            ("Ожидаемый механизм", h.get("mechanism_of_influence", "")),
            ("Целевой KPI", ev.get("kpi_text", "") or "—"),
            ("Объект", f"поток «{t.get('stream', '')}», классы {', '.join(t.get('size_classes') or [])}"),
        ]:
            row = tab.add_row().cells
            row[0].text = k
            row[1].text = str(v)

        d.add_heading("Методика (последовательность экспериментов)", level=2)
        road = h.get("verification_roadmap", [])
        if road:
            for i, v in enumerate(road, 1):
                d.add_paragraph(f"{i}. {v.get('step', '')}", style="List Number")
                d.add_paragraph(f"   Необходимые ресурсы: {v.get('resources', '—')}")
                d.add_paragraph(f"   Критерий успеха/провала: {v.get('success_criteria', '—')}")
        else:
            d.add_paragraph("Определяется исполнителем по согласованию с технологом.")

        d.add_heading("Результаты (заполняется исполнителем)", level=2)
        res = d.add_table(rows=0, cols=2)
        res.style = "Light Grid Accent 1"
        for k in ["Дата проведения", "Исполнитель", "Проба / оборудование",
                  "Результаты измерений", "Вывод (подтверждена / отклонена / доработать)",
                  "Подпись"]:
            row = res.add_row().cells
            row[0].text = k
            row[1].text = ""

        if n < len(hyps) - 1:
            d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
