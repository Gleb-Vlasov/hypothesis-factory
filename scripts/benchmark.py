# -*- coding: utf-8 -*-
"""Бенчмарк против эталонных экспертных гипотез (4 пары организаторов).

Для каждого примера: генерируем гипотезы (LLM или шаблон) и считаем покрытие
экспертных гипотез — по пересечению значимых терминов (Jaccard) и по совпадению
категории механизма. Результат — таблица покрытия + список новых гипотез.
"""
import glob
import io
import json
import os
import re
import sys
import time

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from hypofactory.ingestion.tailings import parse_tailings_report
from hypofactory.rag.retriever import build_default_retriever
from hypofactory.rag.bm25_index import tokenize
from hypofactory.generation.generator import generate_hypotheses
from hypofactory.llm.client import LLMClient
import docx

DATA = os.environ.get(
    "DATA_DIR",
    os.path.join(os.path.dirname(__file__), "..", "..", "DATA"),
)
INDEX = os.path.join(os.path.dirname(__file__), "..", "data_index")
OUT = os.path.join(os.path.dirname(__file__), "..", "data_index", "benchmark_results.json")

# Синонимы/стемминг для доменных терминов: приводим к каноническим меткам тем.
TOPIC_PATTERNS = {
    "измельчение": r"измельч|мельниц|футеров|шаров|доизмельч",
    "классификация": r"классифік|классифик|гидроциклон|насадк|грохот|грохочен|сепарац",
    "флотация": r"флотац|реагент|собират|контактн|агитац|аэрац|пептизат|фронт флотации",
    "плотность": r"плотност|пульп",
    "крупность": r"крупност|гранулометр|тонк|шлам",
    "контроль": r"контрол|автоматизац|регулирован|донастройк",
}


def topics(text: str) -> set:
    low = text.lower()
    return {t for t, pat in TOPIC_PATTERNS.items() if re.search(pat, low)}


def jaccard(a: str, b: str) -> float:
    ta, tb = set(tokenize(a)), set(tokenize(b))
    return len(ta & tb) / len(ta | tb) if ta and tb else 0.0


def expert_hypotheses(folder: str) -> list:
    fs = glob.glob(os.path.join(folder, "Гипотезы*.docx"))
    out = []
    for f in fs:
        d = docx.Document(f)
        lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for row in t.rows:
                for c in row.cells:
                    for para in c.text.split("\n"):
                        if para.strip():
                            lines.append(para.strip())
        for ln in lines:
            if len(ln) > 12 and ln[0].isdigit():
                out.append(re.sub(r"^\d+[.)]\s*", "", ln))
    return list(dict.fromkeys(out))


def match_score(expert: str, ours: list) -> tuple:
    """Лучшее совпадение: комбинируем Jaccard и пересечение тем."""
    et = topics(expert)
    best, best_h = 0.0, None
    for h in ours:
        text = f"{h['title']} {h['statement']}"
        j = jaccard(expert, text)
        t = len(et & topics(text)) / len(et) if et else 0.0
        score = 0.5 * j * 4 + 0.5 * t          # j масштабируем (~0.25 — уже сильное)
        if score > best:
            best, best_h = score, h
    return best, best_h


def main():
    use_llm = os.environ.get("BENCH_LLM", "1") == "1"
    llm = LLMClient() if use_llm else None
    retr = build_default_retriever(INDEX)
    results = []
    files = sorted(glob.glob(os.path.join(DATA, "Пример *")))
    for folder in files:
        xls = glob.glob(os.path.join(folder, "Хвосты*.xlsx"))
        if not xls:
            continue
        name = os.path.basename(xls[0])
        print(f"\n{'='*72}\n{name}")
        t0 = time.time()
        hset = generate_hypotheses(parse_tailings_report(xls[0]), retr, llm=llm)
        ours = [h.to_dict() for h in hset.hypotheses]
        experts = expert_hypotheses(folder)
        covered = 0
        rows = []
        for e in experts:
            score, bh = match_score(e, ours)
            ok = score >= 0.5
            covered += ok
            rows.append({"expert": e, "score": round(score, 2), "matched": ok,
                         "our": bh["title"] if bh else None})
            mark = "✓" if ok else "✗"
            print(f"  {mark} [{score:.2f}] {e[:70]}")
            if bh:
                print(f"        ↳ {bh['title'][:70]}")
        cov = covered / len(experts) if experts else 0
        print(f"  ПОКРЫТИЕ: {covered}/{len(experts)} = {cov:.0%} | наших гипотез: {len(ours)} | {time.time()-t0:.0f}с | режим: {hset.meta['mode']}")
        results.append({"file": name, "coverage": cov, "covered": covered,
                        "experts_total": len(experts), "ours_total": len(ours),
                        "mode": hset.meta["mode"], "rows": rows,
                        "hypotheses": ours})
    json.dump(results, open(OUT, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    total_cov = sum(r["covered"] for r in results) / max(sum(r["experts_total"] for r in results), 1)
    print(f"\n{'='*72}\nИТОГО ПОКРЫТИЕ ЭКСПЕРТНЫХ ГИПОТЕЗ: {total_cov:.0%}  → {OUT}")


if __name__ == "__main__":
    main()
