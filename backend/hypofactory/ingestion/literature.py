"""Загрузка пользовательской литературы в базу знаний (PDF/DOCX/TXT/MD).

Файл режется на пассажи тем же конвейером, что и вшитые учебники
(ingestion/pdf.py), и получает type='user' — такие записи участвуют в поиске
наравне с учебниками и цитируются в гипотезах и диалоге. Идентификаторы
детерминированы (имя файла + номер чанка), поэтому повторная загрузка того же
файла не создаёт дубликатов.
"""
from __future__ import annotations

import os

from hypofactory.ingestion.pdf import chunks_from_pdf, clean_text, chunk_page, is_quality

ALLOWED_EXT = (".pdf", ".docx", ".txt", ".md")
MAX_FILE_MB = 25


def _docx_text(path: str) -> str:
    import docx
    d = docx.Document(path)
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.text.split("\n"):
                    if para.strip():
                        lines.append(para.strip())
    return "\n".join(lines)


def _txt_text(path: str) -> str:
    # типичные кодировки русскоязычных текстов
    for enc in ("utf-8", "cp1251"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()


def parse_document(path: str, original_name: str, meta: dict | None = None) -> list[dict]:
    """Файл → записи корпуса type='user'. Пустой список = текста не извлечено.

    meta — необязательные метаданные источника: author, year, note (условия
    экспериментов и т.п.). Автор/год включаются в отображаемое название,
    поэтому автоматически попадают в цитирование гипотез и диалога.
    """
    ext = os.path.splitext(original_name)[1].lower()
    title = os.path.splitext(os.path.basename(original_name))[0]
    meta = {k: str(v).strip() for k, v in (meta or {}).items() if str(v or "").strip()}
    byline = ", ".join(filter(None, [meta.get("author"), meta.get("year")]))
    if byline:
        title = f"{title} ({byline})"
    from datetime import date
    extra = {**meta, "added_at": date.today().isoformat()}
    records: list[dict] = []

    if ext == ".pdf":
        for i, ch in enumerate(chunks_from_pdf(path, source=original_name, title=title)):
            records.append({
                "id": f"user::{original_name}#p{ch.page}#c{i}",
                "source": original_name, "title": title,
                "type": "user", "page": ch.page, "text": ch.text, **extra,
            })
        return records

    if ext == ".docx":
        text = _docx_text(path)
    elif ext in (".txt", ".md"):
        text = _txt_text(path)
    else:
        raise ValueError(f"Неподдерживаемый формат: {ext} (ожидается PDF/DOCX/TXT/MD)")

    i = 0
    for c in chunk_page(clean_text(text)):
        if is_quality(c):
            records.append({
                "id": f"user::{original_name}#c{i}",
                "source": original_name, "title": title,
                "type": "user", "page": None, "text": c, **extra,
            })
            i += 1
    return records
