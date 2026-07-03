"""Сборка поисковой базы знаний из предоставленных материалов.

Источники:
  - учебники (PDF с текстовым слоем) → пассажи type='textbook';
  - эталонные гипотезы (Гипотезы*.docx) → записи type='hypothesis'
    (одновременно корпус «известных решений» для оценки новизны и few-shot);
  - гайд «Как читать отчёт…» → type='guide'.

Скан без текстового слоя (Лодейщиков) пропускается — off-domain и требует OCR.
"""
from __future__ import annotations

import glob
import os

import docx

from hypofactory.ingestion.pdf import chunks_from_pdf

# Человекочитаемые названия учебников по подстроке имени файла.
BOOK_TITLES = {
    "flotacionnye": "Флотационные методы обогащения",
    "metallurgiya-blagorodnyh": "Металлургия благородных металлов",
    "tehnologiyaobogashcheniyapoleznyh": "Технология обогащения полезных ископаемых",
    "izvlecheniya_zolota_i_serebra": "Извлечение золота и серебра из упорного сырья (статья)",
}
SKIP_PDF = ("lodeyshchikov",)  # скан без текстового слоя


def _title_for(fname: str) -> str:
    low = fname.lower()
    for key, title in BOOK_TITLES.items():
        if key in low:
            return title
    return os.path.splitext(fname)[0]


def _docx_lines(path: str) -> list[str]:
    d = docx.Document(path)
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.text.split("\n"):
                    if para.strip():
                        lines.append(para.strip())
    return lines


def build_corpus(data_dir: str) -> list[dict]:
    records: list[dict] = []
    warnings: list[str] = []

    # --- Учебники (PDF) ---
    pdfs = sorted(glob.glob(os.path.join(data_dir, "**", "*.pdf"), recursive=True))
    for path in pdfs:
        fname = os.path.basename(path)
        if any(s in fname.lower() for s in SKIP_PDF):
            warnings.append(f"PDF пропущен (скан/off-domain): {fname}")
            continue
        title = _title_for(fname)
        chunks = chunks_from_pdf(path, source=fname, title=title)
        if len(chunks) < 5:
            warnings.append(f"PDF мало текста, пропуск: {fname} ({len(chunks)} чанков)")
            continue
        for i, ch in enumerate(chunks):
            records.append({
                "id": f"{fname}#p{ch.page}#c{i}",
                "source": ch.source, "title": ch.title,
                "type": "textbook", "page": ch.page, "text": ch.text,
            })

    # --- Эталонные гипотезы (docx) + гайд ---
    docxs = sorted(glob.glob(os.path.join(data_dir, "**", "*.docx"), recursive=True))
    for path in docxs:
        fname = os.path.basename(path)
        low = fname.lower()
        if low.startswith("~$"):
            continue
        lines = _docx_lines(path)
        if "гипотез" in low:
            fab = fname.replace("Гипотезы", "").replace(".docx", "").strip(" _")
            for j, ln in enumerate(lines):
                # только содержательные строки-гипотезы (нумерованные пункты)
                if len(ln) > 12 and ln[0].isdigit():
                    records.append({
                        "id": f"hyp::{fab}::{j}", "source": fname,
                        "title": f"Экспертные гипотезы: {fab}",
                        "type": "hypothesis", "page": None, "text": ln,
                    })
        elif "читать отчет" in low or "читать отчёт" in low:
            blob = " ".join(lines)
            for k in range(0, len(blob), 700):
                seg = blob[k:k + 800]
                if len(seg) > 80:
                    records.append({
                        "id": f"guide::{k}", "source": fname,
                        "title": "Гайд: как читать отчёт по хвостам",
                        "type": "guide", "page": None, "text": seg,
                    })

    return records, warnings
